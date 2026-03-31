import os
import json
from pathlib import Path
from typing import List, Dict

import docx
from dotenv import load_dotenv
load_dotenv()
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import ValidationError

from ..models.schemas import AdvisorProfile, AdvisorDocument

class DocumentParser:
    def __init__(self, raw_dir: str = "data/raw", processed_dir: str = "data/processed"):
        self.raw_dir = Path(raw_dir)
        self.processed_dir = Path(processed_dir)
        self.processed_dir.mkdir(parents=True, exist_ok=True)
        # We use gpt-4o as it supports strong structured outputs
        self.llm = ChatOpenAI(model="gpt-4o", temperature=0)
        self.parser = PydanticOutputParser(pydantic_object=AdvisorProfile)

    def read_docx(self, file_path: Path) -> str:
        """Extracts text from a docx file."""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    def extract_metadata(self, text: str, filename: str) -> AdvisorProfile:
        """Uses LLM to extract structured metadata (expertise, clients, communication)."""
        prompt = PromptTemplate(
            template=(
                "你是一位專業的理財顧問剖析專家。\n"
                "請根據以下文本內容，提取理專的個人背景資料並符合要求的 JSON 格式。\n"
                "如果某個欄位在文本中未提及或無法推論，請填寫「未提供」。\n"
                "唯一識別碼 (advisor_id) 必須使用以下提供之檔名：{filename}\n\n"
                "{format_instructions}\n\n"
                "文本內容：\n{text}\n"
            ),
            input_variables=["text", "filename"],
            partial_variables={"format_instructions": self.parser.get_format_instructions()}
        )
        
        chain = prompt | self.llm | self.parser
        
        try:
            profile = chain.invoke({"text": text, "filename": filename})
            return profile
        except Exception as e:
            print(f"Failed to parse metadata for {filename}: {e}")
            # Fallback default
            return AdvisorProfile(
                advisor_id=filename,
                name=filename.replace(".docx", ""),
                expertise=["Not Provided"],
                target_clients=["Not Provided"],
                communication_style="Not Provided"
            )

    def process_all(self) -> List[AdvisorDocument]:
        """Reads all docx in raw_dir, extracts JSON metadata, and saves."""
        results = []
        for file_path in self.raw_dir.glob("*.docx"):
            if not file_path.is_file() or file_path.name.startswith("~$"):
                continue
                
            print(f"Processing {file_path.name}...")
            raw_text = self.read_docx(file_path)
            profile = self.extract_metadata(raw_text, file_path.name)
            
            doc = AdvisorDocument(
                profile=profile,
                full_text=raw_text
            )
            results.append(doc)
            
            # Save to processed JSON
            out_file = self.processed_dir / f"{file_path.stem}.json"
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(doc.model_dump_json(indent=2))
                
        return results

if __name__ == "__main__":
    parser = DocumentParser()
    docs = parser.process_all()
    print(f"Processed {len(docs)} documents.")
